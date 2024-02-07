import openai
import threading
import os
from gtts import gTTS
from pygame import mixer
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Set your OpenAI API key
openai.api_key = "AIzaSyDjG3rU6dhRxr_m1V6_TsITSYgJVWl06L8"

# Initialize the Pygame mixer for audio playback
mixer.init()

def generate_response(prompt):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message['content']
    except Exception as e:
        return str(e)

def save_text_and_speech(text, filename):
    print(text)

    # Save the text to a text file
    with open(filename + ".txt", 'w', encoding='utf-8') as file:
        file.write(text)

    # Convert text to speech with gTTS and handle errors
    try:
        tts = gTTS(text=text, lang='en', slow=False)
        tts.save(filename + ".mp3")

        # Play the generated audio
        mixer.music.load(filename + ".mp3")
        mixer.music.play()
        while mixer.music.get_busy():
            pass  # Wait for audio to finish playing
    except Exception as e:
        print("Error during audio generation or playback:", str(e))

def create_word_document(filename, responses):
    doc = Document()

    # Add a title to the Word document
    title = doc.add_heading('Generated Responses', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for idx, response in enumerate(responses, 1):
        # Add a heading for each response
        heading = doc.add_heading(f'Response {idx}', level=2)

        # Add the response text as a paragraph
        paragraph = doc.add_paragraph(response)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Set font size for response text
        for run in paragraph.runs:
            font = run.font
            font.size = Pt(12)

        # Add a page break between responses (except for the last one)
        if idx < len(responses):
            run = paragraph.add_run()
            run.add_break()

    # Save the Word document
    doc.save(filename + ".docx")

def main():
    counter = 1
    responses = []  # To store responses for the Word document

    while True:
        # Display a menu of options to the user
        print("Choose an option:")
        print("1. Black Box Testing")
        print("2. White Box Testing")
        print("3. Unit Testing")
        print("4. Quit")

        # Get the user's choice
        choice = input("Enter your choice (1/2/3/4): ")

        if choice == "1":
            topic = "Black Box Testing"
        elif choice == "2":
            topic = "White Box Testing"
        elif choice == "3":
            topic = "Unit Testing"
        elif choice == "4":
            print("Goodbye!")
            break
        else:
            print("Invalid choice. Please choose a valid option.")
            continue

        # Create a prompt with the chosen topic
        prompt = f"Please provide an explanation of {topic} in a conversational and friendly tone suitable for a trainer to narrate to students."

        # Generate Response
        response = generate_response(prompt)

        # Add the response to the list of responses
        responses.append(response)

        # Save the response as both text and speech (MP3)
        filename = f"response_{counter}"
        save_text_and_speech(response, filename)

        counter += 1

        # Sleep for a while (e.g., 10 seconds) before the next iteration
        time.sleep(10)  # Adjust the sleep duration as needed

    # Create a Word document with the collected responses
    create_word_document("generated_responses", responses)

if __name__ == "__main__":
    main()
